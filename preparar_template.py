# -*- coding: utf-8 -*-
"""
Etapa de preparação (rodar uma única vez): converte o modelo original
"1. Requerimento (GERID) - Atualizado.docx" (com trechos grifados em amarelo
como placeholder) em modelos/template.docx, com tokens Jinja2 no lugar dos
trechos grifados e sem nenhum realce amarelo.

Não usa substituição "às cegas" de texto: percorre os runs do documento
em ordem, exige que cada run com realce amarelo tenha exatamente o texto
já conhecido do modelo (levantado manualmente a partir do documento
original) e só então troca o conteúdo. Se o modelo for editado e os runs
não corresponderem mais ao esperado, o script para com um erro claro em
vez de gerar um template.docx incorreto silenciosamente.
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

PASTA_BASE = Path(__file__).resolve().parent
ORIGEM_PADRAO = PASTA_BASE / "modelos" / "1. Requerimento (GERID) - Atualizado.docx"
DESTINO_PADRAO = PASTA_BASE / "modelos" / "template.docx"

# Sequência exata (na ordem em que aparecem no documento) de cada run com
# w:highlight="yellow". `None` em "novo" significa "manter o texto, só
# remover o realce" (texto fixo do modelo que por algum motivo também
# estava grifado, mas não é um placeholder de dados). Vários runs aqui
# misturam texto fixo com o placeholder de "X" na mesma run (o autor do
# modelo não separou em runs distintas como antes) - por isso o texto novo
# às vezes troca só um trecho da run, mantendo o restante literal.
RUNS_GRIFADOS = [
    # Razão social da empresa representada
    ("xxxxxxxxxxxxxxxxx", "{{ empresa }}"),
    # CNPJ da empresa representada (campo novo nesta versão do modelo)
    ("xxxxxxxxxxxxxxxxx", "{{ cnpj }}"),
    # "O segurado XXXX, inscrito no CPF nº XXXX e sob o NIT nº XXXX, era empregado"
    ("O segurado XXXXXXXXXXXXXX", "{{ tratamento_cap }} {{ segurado }}"),
    (", inscrito no CPF nº XXXXXXXXXX e sob o NIT nº ", ", {{ inscrito }} no CPF nº {{ cpf }} e sob o NIT nº "),
    ("XXXXXXXXXXXXX, era empregado", "{{ nit }}, era {{ empregado }}"),
    # "o benefício XXXX, da espécie B9X, nº XXXX"
    ("o ", None),
    ("benefício XXXXXXXXXXX", "benefício de {{ beneficio }}"),
    (", da ", None),
    ("espécie B9X", "espécie {{ especie }}"),
    (",", None),
    (" nº XXXXXXXXXXXX", " nº {{ nb }}"),
]


def cor_realce(run):
    """Retorna o valor de w:highlight do run, ou None se não houver realce."""
    rpr = run._element.rPr
    if rpr is None:
        return None
    highlight = rpr.find(qn("w:highlight"))
    if highlight is None:
        return None
    return highlight.get(qn("w:val"))


def remover_realce(run):
    run.font.highlight_color = None


def runs_do_corpo(doc):
    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            yield run


def aplicar_token_data(doc):
    """Substitui o campo de data automático do Word ("Florianópolis, {TIME...}.")
    por um único run de texto "{{ data_extenso }}", controlável pelo gerar.py.
    """
    alvo = None
    for paragrafo in doc.paragraphs:
        if paragrafo.runs and paragrafo.runs[0].text.strip().startswith("Florianópolis"):
            alvo = paragrafo
            break
    if alvo is None:
        raise RuntimeError(
            'Não encontrei o parágrafo "Florianópolis, ..." para tratar a data do rodapé.'
        )

    runs = alvo.runs
    # Localiza o run literal "Florianópolis, " e o run final "." que fecham o campo de data.
    idx_inicio = next(
        (i for i, r in enumerate(runs) if r.text.strip().startswith("Florianópolis")), None
    )
    if idx_inicio is None:
        raise RuntimeError('Run "Florianópolis, " não encontrado como esperado.')

    # Os runs entre o início (exclusive) e o run "." final são o mecanismo de campo
    # (fldChar begin/instrText/fldChar separate/texto em cache/fldChar end).
    idx_fim = None
    for i in range(idx_inicio + 1, len(runs)):
        if runs[i].text == ".":
            idx_fim = i
            break
    if idx_fim is None:
        raise RuntimeError('Run "." de fechamento da data não encontrado como esperado.')

    runs_campo = runs[idx_inicio + 1 : idx_fim]
    if not runs_campo:
        raise RuntimeError("Estrutura do campo de data não corresponde ao esperado.")

    # Usa o primeiro run do campo para herdar a formatação e nele coloca o token;
    # remove os demais runs do campo (fldChar/instrText/texto em cache).
    primeiro = runs_campo[0]
    primeiro.text = "{{ data_extenso }}"
    for extra in runs_campo[1:]:
        extra._element.getparent().remove(extra._element)


def preparar(origem: Path, destino: Path) -> None:
    if not origem.exists():
        raise FileNotFoundError(f"Documento original não encontrado: {origem}")

    doc = Document(str(origem))

    grifados = [run for run in runs_do_corpo(doc) if cor_realce(run) == "yellow"]

    if len(grifados) != len(RUNS_GRIFADOS):
        raise RuntimeError(
            f"Esperava {len(RUNS_GRIFADOS)} runs grifados em amarelo, mas encontrei "
            f"{len(grifados)} no documento de origem. O modelo pode ter sido editado; "
            "revise a lista RUNS_GRIFADOS em preparar_template.py antes de continuar."
        )

    for i, (run, (esperado, novo)) in enumerate(zip(grifados, RUNS_GRIFADOS), start=1):
        if run.text != esperado:
            raise RuntimeError(
                f"Run grifado nº {i}: texto era {run.text!r}, esperava {esperado!r}. "
                "O modelo pode ter sido editado; revise RUNS_GRIFADOS antes de continuar."
            )
        if novo is not None:
            run.text = novo
        remover_realce(run)

    aplicar_token_data(doc)

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    print(f"Template gerado em: {destino}")


def main():
    parser = argparse.ArgumentParser(
        description="Gera modelos/template.docx a partir do Requerimento - GERID original."
    )
    parser.add_argument("--origem", type=Path, default=ORIGEM_PADRAO, help="Caminho do .docx original")
    parser.add_argument("--destino", type=Path, default=DESTINO_PADRAO, help="Caminho de saída do template")
    args = parser.parse_args()

    try:
        preparar(args.origem, args.destino)
    except (FileNotFoundError, RuntimeError) as exc:
        print(f"ERRO: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
