# -*- coding: utf-8 -*-
"""
Etapa de preparação (rodar uma única vez): converte o modelo original
"Modelo Exigência.docx" (com trechos "XXXX" digitados manualmente como
placeholder, sem realce de cor) em modelos/template_exigencia.docx, com
tokens Jinja2 no lugar desses trechos.

Diferente de preparar_template.py, este modelo não usa realce amarelo nos
placeholders — por isso a localização é feita por (parágrafo, run), com
verificação do texto exato esperado em cada um antes de substituir. Se o
modelo for editado e os runs não corresponderem mais ao esperado, o script
para com um erro claro em vez de gerar um template incorreto silenciosamente.
"""

import argparse
import sys
from pathlib import Path

from docx import Document

PASTA_BASE = Path(__file__).resolve().parent
ORIGEM_PADRAO = PASTA_BASE / "modelos" / "Modelo Exigência.docx"
DESTINO_PADRAO = PASTA_BASE / "modelos" / "template_exigencia.docx"

# Cada item: (índice do parágrafo, índice do run dentro do parágrafo,
# texto exato esperado, texto novo).
SUBSTITUICOES = [
    # Parág. 5: "...para representar a empresa XXXX, ... inscrita no CNPJ sob o nº XXXX..."
    (5, 1, "XXXXXXXXXXXXXXXXXX", "{{ empresa }}"),
    (5, 3, "XXXXXXXXXXXXXXX", "{{ cnpj }}"),
    # Parág. 35: "O segurado XXXX, inscrito no CPF nº XXXX e sob o NIT nº XXXX,
    # era empregado ... benefício de auxílio ... da espécie B91, nº XXXX"
    (35, 0, "O segurado ", "{{ tratamento_cap }} "),
    (35, 1, "XXXXXXXXXXXXX", "{{ segurado }}"),
    (35, 2, ", inscrito no CPF nº ", ", {{ inscrito }} no CPF nº "),
    (35, 3, "XXXXXXXXX", "{{ cpf }}"),
    (35, 5, "XXXXXXXXXXXXXXX", "{{ nit }}"),
    (
        35, 6,
        ", era empregado da empresa representada quando lhe foi concedido o ",
        ", era {{ empregado }} da empresa representada quando lhe foi concedido o ",
    ),
    (
        35, 7,
        "benefício de auxílio por incapacidade temporária por acidente de trabalho, da espécie B91, nº ",
        "benefício de {{ beneficio }}, da espécie {{ especie }}, nº ",
    ),
    (35, 8, "XXXXXXXXXXXXX", "{{ nb }}"),
    # Parág. 39: "...comprovação da representatividade da empresa xxxx..."
    (39, 3, "xxxxxxxxxxxxxx", "{{ empresa }}"),
    # Parág. 45: data do rodapé, já como texto literal (sem campo automático do Word)
    (45, 0, "Florianópolis, 24 de julho de 2026.", "Florianópolis, {{ data_extenso }}."),
]


def preparar(origem: Path, destino: Path) -> None:
    if not origem.exists():
        raise FileNotFoundError(f"Documento original não encontrado: {origem}")

    doc = Document(str(origem))

    for indice_paragrafo, indice_run, esperado, novo in SUBSTITUICOES:
        if indice_paragrafo >= len(doc.paragraphs):
            raise RuntimeError(
                f"Parágrafo {indice_paragrafo} não existe no documento "
                f"(documento tem {len(doc.paragraphs)} parágrafos). O modelo pode ter sido editado."
            )
        paragrafo = doc.paragraphs[indice_paragrafo]
        if indice_run >= len(paragrafo.runs):
            raise RuntimeError(
                f"Parágrafo {indice_paragrafo}: run {indice_run} não existe "
                f"(parágrafo tem {len(paragrafo.runs)} runs). O modelo pode ter sido editado."
            )
        run = paragrafo.runs[indice_run]
        if run.text != esperado:
            raise RuntimeError(
                f"Parágrafo {indice_paragrafo}, run {indice_run}: texto era {run.text!r}, "
                f"esperava {esperado!r}. O modelo pode ter sido editado; revise SUBSTITUICOES "
                "antes de continuar."
            )
        run.text = novo

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    print(f"Template gerado em: {destino}")


def main():
    parser = argparse.ArgumentParser(
        description="Gera modelos/template_exigencia.docx a partir do Modelo Exigência.docx original."
    )
    parser.add_argument("--origem", type=Path, default=ORIGEM_PADRAO, help="Caminho do .docx original")
    parser.add_argument("--destino", type=Path, default=DESTINO_PADRAO, help="Caminho de saída do template")
    args = parser.parse_args()

    try:
        preparar(args.origem, args.destino)
    except (FileNotFoundError, RuntimeError) as exc:
        print(f"ERRO: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
